import os
import shutil
from datetime import datetime, timezone
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Query, WebSocket, WebSocketDisconnect, Depends, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from config import UPLOAD_DIR
from meeting_runner import process_meeting, parse_meeting_output
from db import meetings, conversations, users
from chat_agents import chat_agent
from meeting_agents import master_agent, live_points_agent
from tools import live_transcribe_chunk, transcribe_audio, translate_text, optimize_text, generate_meeting_notes
from agents import Runner
import uuid
import json
import io
from pydub import AudioSegment
from docx import Document
import csv
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import firebase_admin
from firebase_admin import credentials, auth
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Firebase Admin SDK
FIREBASE_CREDENTIALS_PATH = os.getenv("FIREBASE_CREDENTIALS_PATH", "serviceAccountKey.json")
if not firebase_admin._apps:
    try:
        cred = credentials.Certificate(FIREBASE_CREDENTIALS_PATH)
        firebase_admin.initialize_app(cred)
    except FileNotFoundError:
        raise Exception(f"Firebase credentials file not found at {FIREBASE_CREDENTIALS_PATH}")

# Initialize FastAPI app
app = FastAPI(title="TalkToText Pro API")

# Allow CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "ws://localhost:3000", "ws://localhost:3001"],
    allow_credentials=True,
    allow_methods=["*", "Authorization"],
    allow_headers=["Authorization"],
)

# Ensure upload directory exists
os.makedirs(UPLOAD_DIR, exist_ok=True)
security = HTTPBearer()

# Dependency to get authenticated user ID
async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        decoded_token = auth.verify_id_token(token)
        return decoded_token["uid"]
    except Exception as e:
        raise HTTPException(status_code=401, detail="Invalid or expired token")

# Live Session Manager for WebSocket
class LiveSessionManager:
    def __init__(self):
        self.sessions = {}  # run_id -> {'websocket': WebSocket, 'transcript': '', 'accumulated_audio': b'', 'language': str, 'user_id': str}

    async def handle_live_session(self, websocket: WebSocket, run_id: str, language: str, user_id: str):
        self.sessions[run_id] = {
            'websocket': websocket,
            'transcript': '',
            'accumulated_audio': b'',
            'language': language,
            'user_id': user_id
        }
        await websocket.accept()
        logger.info(f"WebSocket session started for run_id: {run_id}, user_id: {user_id}")

        try:
            while True:
                data = await websocket.receive()
                if data["type"] == "websocket.disconnect":
                    logger.info(f"WebSocket disconnected for run_id: {run_id}")
                    break
                elif data["type"] == "websocket.receive":
                    if "bytes" in data:
                        session = self.sessions[run_id]
                        session['accumulated_audio'] += data["bytes"]
                        logger.info(f"Received audio chunk for run_id: {run_id}, size: {len(data['bytes'])}")

                        if len(session['accumulated_audio']) >= 1024:  # Process smaller chunks
                            chunk_transcript = await live_transcribe_chunk(session['accumulated_audio'], session['language'])
                            if chunk_transcript and not chunk_transcript.startswith("Error"):
                                optimized = await optimize_text(chunk_transcript)
                                if session['language'] != 'en':
                                    optimized = await translate_text(optimized, 'en')
                                session['transcript'] += optimized + ' '
                                logger.info(f"Processed transcript for run_id: {run_id}: {optimized}")
                            else:
                                logger.warning(f"No valid transcript for chunk in run_id: {run_id}")

                            runner = Runner()
                            context = {"transcript": session['transcript']}
                            result = await runner.run(
                                live_points_agent,
                                input="Generate updated key points and action items from the live transcript.",
                                context=context
                            )
                            points = parse_meeting_output(result.final_output)

                            await meetings.update_one(
                                {"run_id": run_id},
                                {"$set": {"partial_transcript": session['transcript'], "partial_points": points}}
                            )

                            await websocket.send_json({
                                "run_id": run_id,
                                "transcript": optimized if chunk_transcript and not chunk_transcript.startswith("Error") else "",
                                "key_points": points.get("key_points", []),
                                "action_items": points.get("action_items", []),
                                "sentiment": points.get("sentiment", "neutral")
                            })
                            logger.info(f"Sent transcription update for run_id: {run_id}")

                            session['accumulated_audio'] = b''
                    elif "text" in data:
                        try:
                            message = json.loads(data["text"])
                            if message.get("type") == "ping":
                                logger.info(f"Received ping for run_id: {run_id}")
                                await websocket.send_json({"type": "pong"})
                                continue
                        except json.JSONDecodeError:
                            logger.error(f"Invalid JSON in WebSocket text message for run_id: {run_id}")
        except WebSocketDisconnect:
            logger.info(f"WebSocket disconnected for run_id: {run_id}")
        except Exception as e:
            logger.error(f"Error in live session for run_id {run_id}: {str(e)}")
        finally:
            if run_id in self.sessions:
                transcript = self.sessions[run_id]['transcript']
                del self.sessions[run_id]
                await meetings.update_one(
                    {"run_id": run_id},
                    {"$set": {"status": "done", "result": parse_meeting_output(transcript)}}
                )
                logger.info(f"Session cleaned up for run_id: {run_id}")

live_manager = LiveSessionManager()

@app.websocket("/live-transcribe")
async def live_transcribe(
    websocket: WebSocket,
    language: str = Query("en"),
    token: str = Query(None),
    user_id: str = Query(None)
):
    try:
        # Verify the token
        decoded_token = auth.verify_id_token(token)
        if decoded_token["uid"] != user_id:
            logger.error(f"User ID mismatch: token UID {decoded_token['uid']}, provided user_id {user_id}")
            await websocket.close(code=1008, reason="User ID mismatch")
            return

        run_id = str(uuid.uuid4())
        await meetings.insert_one({
            "run_id": run_id,
            "user_id": user_id,
            "status": "live",
            "file_path": None,
            "language": language,
            "filename": f"Live Meeting {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}",
            "created_at": datetime.now(timezone.utc)
        })
        logger.info(f"Starting live transcription for run_id: {run_id}, user_id: {user_id}")
        await live_manager.handle_live_session(websocket, run_id, language, user_id)
    except Exception as e:
        logger.error(f"Authentication or initialization failed: {str(e)}")
        await websocket.close(code=1008, reason=f"Authentication failed: {str(e)}")

# [Rest of the server.py code remains unchanged]
@app.post("/upload-audio")
async def upload_audio(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user),
    language: str = Query("en", description="Language of the meeting audio")
):
    valid_extensions = (".mp3", ".wav", ".m4a", ".mp4")
    if file.content_type.split("/")[0] not in ["audio", "video"] and not file.filename.endswith(valid_extensions):
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload audio/video only.")
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    unique_filename = f"{user_id}_{timestamp}_{file.filename}"
    saved_path = os.path.join(UPLOAD_DIR, unique_filename)
    try:
        with open(saved_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error saving file: {str(e)}")
    run_id, _ = await process_meeting(saved_path, user_id, language)
    background_tasks.add_task(process_meeting, saved_path, user_id, language)
    return {
        "status": "processing",
        "run_id": run_id,
        "message": "File uploaded successfully and processing started.",
        "file_info": {
            "filename": file.filename,
            "saved_as": unique_filename,
            "content_type": file.content_type,
            "path": saved_path,
            "language": language
        }
    }

@app.post("/chat/{run_id}")
async def chat_with_meeting(
    run_id: str,
    request: dict,
    user_id: str = Depends(get_current_user)
):
    message = request.get("message", "")
    if not message:
        raise HTTPException(400, "Message is required")
  
    meeting = await meetings.find_one({"run_id": run_id, "user_id": user_id})
    if not meeting:
        raise HTTPException(404, "Meeting not found")
  
    if meeting.get("status") != "done":
        raise HTTPException(400, "Meeting processing not complete yet")
  
    context = {
        "meeting_transcript": meeting.get("result", {}).get("final_output", ""),
        "meeting_summary": meeting.get("result", {}).get("summary", ""),
        "meeting_key_points": meeting.get("result", {}).get("key_points", []),
        "meeting_decisions": meeting.get("result", {}).get("decisions", []),
        "meeting_action_items": meeting.get("result", {}).get("action_items", []),
        "original_file": meeting.get("file_path", ""),
        "language": meeting.get("language", "en")
    }
  
    runner = Runner()
    result = await runner.run(
        chat_agent,
        input=message,
        context=context
    )
  
    chat_id = str(uuid.uuid4())
    await conversations.insert_one({
        "chat_id": chat_id,
        "run_id": run_id,
        "user_id": user_id,
        "message": message,
        "response": result.final_output,
        "timestamp": datetime.now(timezone.utc)
    })
  
    return {
        "chat_id": chat_id,
        "response": result.final_output,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }

@app.get("/conversations/{run_id}")
async def get_conversations(run_id: str, user_id: str = Depends(get_current_user)):
    cursor = conversations.find({"run_id": run_id, "user_id": user_id}).sort("timestamp", 1)
    chats = await cursor.to_list(length=50)
  
    for chat in chats:
        chat["_id"] = str(chat["_id"])
        if chat.get("timestamp"):
            chat["timestamp"] = chat["timestamp"].isoformat()
  
    return {
        "run_id": run_id,
        "conversations": chats
    }

@app.delete("/conversations/{run_id}/{chat_id}")
async def delete_conversation(run_id: str, chat_id: str, user_id: str = Depends(get_current_user)):
    result = await conversations.delete_one({"chat_id": chat_id, "run_id": run_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Conversation not found")
    return {"message": "Conversation deleted successfully"}

@app.delete("/conversations/{run_id}")
async def delete_all_conversations(run_id: str, user_id: str = Depends(get_current_user)):
    result = await conversations.delete_many({"run_id": run_id, "user_id": user_id})
    return {"message": f"Deleted {result.deleted_count} conversations"}

@app.put("/conversations/{run_id}/{chat_id}")
async def update_conversation(run_id: str, chat_id: str, request: dict, user_id: str = Depends(get_current_user)):
    message = request.get("message")
    response = request.get("response")
    if not message and not response:
        raise HTTPException(400, "At least one of 'message' or 'response' must be provided")
   
    update_data = {}
    if message:
        update_data["message"] = message
    if response:
        update_data["response"] = response
    update_data["updated_at"] = datetime.now(timezone.utc)
   
    result = await conversations.update_one(
        {"chat_id": chat_id, "run_id": run_id, "user_id": user_id},
        {"$set": update_data}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Conversation not found")
    return {"message": "Conversation updated successfully"}

@app.get("/conversations/{run_id}/download/word")
async def download_conversations_word(run_id: str, user_id: str = Depends(get_current_user)):
    meeting = await meetings.find_one({"run_id": run_id, "user_id": user_id})
    if not meeting:
        raise HTTPException(404, "Meeting not found")
   
    cursor = conversations.find({"run_id": run_id, "user_id": user_id}).sort("timestamp", 1)
    chats = await cursor.to_list(length=50)
   
    doc = Document()
    doc.add_heading(f"Meeting: {meeting.get('filename', 'Untitled Meeting')}", 0)
    doc.add_heading("Meeting Details", level=1)
    doc.add_paragraph(f"Date: {meeting.get('created_at', '').isoformat()}")
    doc.add_paragraph(f"Language: {meeting.get('language', 'en')}")
    doc.add_paragraph(f"Status: {meeting.get('status', 'unknown')}")
   
    if meeting.get("result"):
        doc.add_heading("Transcript", level=2)
        doc.add_paragraph(meeting["result"].get("final_output", ""))
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(meeting["result"].get("summary", ""))
        doc.add_heading("Key Points", level=2)
        for point in meeting["result"].get("key_points", []):
            doc.add_paragraph(f"- {point}")
        doc.add_heading("Action Items", level=2)
        for item in meeting["result"].get("action_items", []):
            doc.add_paragraph(f"- {item}")
   
    if chats:
        doc.add_heading("Conversations", level=1)
        for chat in chats:
            doc.add_heading(f"Chat ID: {chat['chat_id']}", level=2)
            doc.add_paragraph(f"Timestamp: {chat['timestamp'].isoformat()}")
            doc.add_paragraph(f"User Message: {chat['message']}")
            doc.add_paragraph(f"AI Response: {chat['response']}")
            doc.add_paragraph("-" * 50)
   
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
   
    filename = f"meeting_{run_id}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/conversations/{run_id}/download/csv")
async def download_conversations_csv(run_id: str, user_id: str = Depends(get_current_user)):
    meeting = await meetings.find_one({"run_id": run_id, "user_id": user_id})
    if not meeting:
        raise HTTPException(404, "Meeting not found")
   
    cursor = conversations.find({"run_id": run_id, "user_id": user_id}).sort("timestamp", 1)
    chats = await cursor.to_list(length=50)
   
    buffer = io.StringIO()
    writer = csv.writer(buffer, quoting=csv.QUOTE_MINIMAL)
    writer.writerow(["Meeting ID", "Meeting Name", "Date", "Language", "Status", "Type", "Content"])
   
    writer.writerow([
        meeting["run_id"],
        meeting.get("filename", "Untitled Meeting"),
        meeting.get("created_at", "").isoformat(),
        meeting.get("language", "en"),
        meeting.get("status", "unknown"),
        "Transcript",
        meeting.get("result", {}).get("final_output", "")
    ])
    writer.writerow(["", "", "", "", "", "Summary", meeting.get("result", {}).get("summary", "")])
    for point in meeting.get("result", {}).get("key_points", []):
        writer.writerow(["", "", "", "", "", "Key Point", point])
    for item in meeting.get("result", {}).get("action_items", []):
        writer.writerow(["", "", "", "", "", "Action Item", item])
   
    for chat in chats:
        writer.writerow([
            meeting["run_id"],
            meeting.get("filename", "Untitled Meeting"),
            chat["timestamp"].isoformat(),
            "",
            "",
            "User Message",
            chat["message"]
        ])
        writer.writerow([
            meeting["run_id"],
            meeting.get("filename", "Untitled Meeting"),
            chat["timestamp"].isoformat(),
            "",
            "",
            "AI Response",
            chat["response"]
        ])
   
    buffer.seek(0)
    filename = f"meeting_{run_id}.csv"
    return StreamingResponse(
        io.BytesIO(buffer.getvalue().encode('utf-8')),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.delete("/meetings/{run_id}")
async def delete_meeting(run_id: str, user_id: str = Depends(get_current_user)):
    meeting = await meetings.find_one({"run_id": run_id, "user_id": user_id})
    if not meeting:
        raise HTTPException(404, "Meeting not found")
   
    if meeting.get("file_path"):
        try:
            os.remove(meeting["file_path"])
        except OSError:
            pass
   
    await meetings.delete_one({"run_id": run_id, "user_id": user_id})
    await conversations.delete_many({"run_id": run_id, "user_id": user_id})
    return {"message": "Meeting and associated conversations deleted successfully"}

@app.put("/meetings/{run_id}")
async def rename_meeting(run_id: str, request: dict, user_id: str = Depends(get_current_user)):
    filename = request.get("filename")
    if not filename:
        raise HTTPException(400, "Filename is required")
   
    result = await meetings.update_one(
        {"run_id": run_id, "user_id": user_id},
        {"$set": {"filename": filename}}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Meeting not found")
    return {"message": "Meeting renamed successfully"}

@app.get("/meetings/{user_id}")
async def get_all_meetings(user_id: str = Depends(get_current_user)):
    cursor = meetings.find({"user_id": user_id}).sort("created_at", -1)
    meetings_list = await cursor.to_list(length=100)
  
    simplified_meetings = []
    for meeting in meetings_list:
        simplified_meetings.append({
            "run_id": meeting.get("run_id"),
            "filename": meeting.get("filename", meeting.get("file_path", "").split("/")[-1] if meeting.get("file_path") else "Live Meeting"),
            "created_at": meeting.get("created_at").isoformat() if meeting.get("created_at") else "",
            "status": meeting.get("status"),
            "has_result": meeting.get("status") == "done",
            "language": meeting.get("language", "en")
        })
  
    return {
        "user_id": user_id,
        "meetings": simplified_meetings
    }

@app.get("/status/{run_id}")
async def get_status(run_id: str, user_id: str = Depends(get_current_user)):
    doc = await meetings.find_one({"run_id": run_id, "user_id": user_id})
    if not doc:
        raise HTTPException(404, "Run ID not found")
    return {
        "run_id": run_id,
        "status": doc.get("status", "unknown"),
        "result": doc.get("result"),
        "language": doc.get("language", "en")
    }

@app.get("/history/{user_id}")
async def history(user_id: str = Depends(get_current_user)):
    cursor = meetings.find({"user_id": user_id}).sort("created_at", -1).limit(20)
    docs = await cursor.to_list(length=20)
  
    for doc in docs:
        doc["_id"] = str(doc["_id"])
        if doc.get("created_at"):
            doc["created_at"] = doc["created_at"].isoformat()
  
    return {
        "user_id": user_id,
        "history": docs
    }

@app.get("/")
async def root():
    return {"message": "Server is running successfully!"}

@app.get("/ping")
async def ping():
    return {"status": "ok", "message": "TalkToText API is live 🚀"}

@app.get("/test-gemini")
async def test_gemini():
    from meeting_agents import master_agent
    from agents import Runner
    runner = Runner()
    test_result = await runner.run(
        master_agent,
        input="Hello Gemini! Summarize this meeting: 'We discussed project timelines and assigned tasks.'"
    )
    return {"output": test_result.final_output}

@app.post("/auth/save-user")
async def save_user(user: dict = Body(...)):
    user_id = user.get("uid")
    if not user_id:
        raise HTTPException(400, "Missing uid")
    existing_user = await users.find_one({"uid": user_id})
    if existing_user:
        await users.update_one(
            {"uid": user_id},
            {"$set": {
                "email": user.get("email"),
                "displayName": user.get("displayName"),
                "photoURL": user.get("photoURL"),
                "provider": user.get("provider", "email"),
                "last_login": datetime.now(timezone.utc)
            }}
        )
        return {"message": "User updated", "uid": user_id}
    else:
        user["created_at"] = datetime.now(timezone.utc)
        user["last_login"] = datetime.now(timezone.utc)
        await users.insert_one(user)
        return {"message": "User created", "uid": user_id}
