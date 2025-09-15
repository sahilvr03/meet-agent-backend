from fastapi import APIRouter, HTTPException, Depends
from datetime import datetime, timezone
from fastapi.responses import StreamingResponse
from db import meetings, conversations
from agents import Runner
from chat_agents import chat_agent
from docx import Document
import csv
import io
import uuid

router = APIRouter()

@router.post("/{run_id}")
async def chat_with_meeting(
    run_id: str,
    request: dict,
    user_id: str = Depends(get_current_user)
):
    message = request.get("message", "")
    if not message:
        raise HTTPException(400, "Message is required")
  
    meeting = await meetings.find_one({"run_id": run_id, "user_id": user_id})
    if not meeting:
        raise HTTPException(404, "Meeting not found")
  
    if meeting.get("status") != "done":
        raise HTTPException(400, "Meeting processing not complete yet")
  
    context = {
        "meeting_transcript": meeting.get("result", {}).get("final_output", ""),
        "meeting_summary": meeting.get("result", {}).get("summary", ""),
        "meeting_key_points": meeting.get("result", {}).get("key_points", []),
        "meeting_decisions": meeting.get("result", {}).get("decisions", []),
        "meeting_action_items": meeting.get("result", {}).get("action_items", []),
        "original_file": meeting.get("file_path", ""),
        "language": meeting.get("language", "en")
    }
  
    runner = Runner()
    result = await runner.run(
        chat_agent,
        input=message,
        context=context
    )
  
    chat_id = str(uuid.uuid4())
    await conversations.insert_one({
        "chat_id": chat_id,
        "run_id": run_id,
        "user_id": user_id,
        "message": message,
        "response": result.final_output,
        "timestamp": datetime.now(timezone.utc)
    })
  
    return {
        "chat_id": chat_id,
        "response": result.final_output,
        "timestamp": datetime.now(timezone.utc).isoformat()
    }

@router.get("/{run_id}")
async def get_conversations(run_id: str, user_id: str = Depends(get_current_user)):
    cursor = conversations.find({"run_id": run_id, "user_id": user_id}).sort("timestamp", 1)
    chats = await cursor.to_list(length=50)
  
    for chat in chats:
        chat["_id"] = str(chat["_id"])
        if chat.get("timestamp"):
            chat["timestamp"] = chat["timestamp"].isoformat()
  
    return {
        "run_id": run_id,
        "conversations": chats
    }

@router.delete("/{run_id}/{chat_id}")
async def delete_conversation(run_id: str, chat_id: str, user_id: str = Depends(get_current_user)):
    result = await conversations.delete_one({"chat_id": chat_id, "run_id": run_id, "user_id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(404, "Conversation not found")
    return {"message": "Conversation deleted successfully"}

@router.delete("/{run_id}")
async def delete_all_conversations(run_id: str, user_id: str = Depends(get_current_user)):
    result = await conversations.delete_many({"run_id": run_id, "user_id": user_id})
    return {"message": f"Deleted {result.deleted_count} conversations"}

@router.put("/{run_id}/{chat_id}")
async def update_conversation(run_id: str, chat_id: str, request: dict, user_id: str = Depends(get_current_user)):
    message = request.get("message")
    response = request.get("response")
    if not message and not response:
        raise HTTPException(400, "At least one of 'message' or 'response' must be provided")
   
    update_data = {}
    if message:
        update_data["message"] = message
    if response:
        update_data["response"] = response
    update_data["updated_at"] = datetime.now(timezone.utc)
   
    result = await conversations.update_one(
        {"chat_id": chat_id, "run_id": run_id, "user_id": user_id},
        {"$set": update_data}
    )
    if result.matched_count == 0:
        raise HTTPException(404, "Conversation not found")
    return {"message": "Conversation updated successfully"}

@router.get("/{run_id}/download/word")
async def download_conversations_word(run_id: str, user_id: str = Depends(get_current_user)):
    meeting = await meetings.find_one({"run_id": run_id, "user_id": user_id})
    if not meeting:
        raise HTTPException(404, "Meeting not found")
   
    cursor = conversations.find({"run_id": run_id, "user_id": user_id}).sort("timestamp", 1)
    chats = await cursor.to_list(length=50)
   
    doc = Document()
    doc.add_heading(f"Meeting: {meeting.get('filename', 'Untitled Meeting')}", 0)
    doc.add_heading("Meeting Details", level=1)
    doc.add_paragraph(f"Date: {meeting.get('created_at', '').isoformat()}")
    doc.add_paragraph(f"Language: {meeting.get('language', 'en')}")
    doc.add_paragraph(f"Status: {meeting.get('status', 'unknown')}")
   
    if meeting.get("result"):
        doc.add_heading("Transcript", level=2)
        doc.add_paragraph(meeting["result"].get("final_output", ""))
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(meeting["result"].get("summary", ""))
        doc.add_heading("Key Points", level=2)
        for point in meeting["result"].get("key_points", []):
            doc.add_paragraph(f"- {point}")
        doc.add_heading("Action Items", level=2)
        for item in meeting["result"].get("action_items", []):
            doc.add_paragraph(f"- {item}")
   
    if chats:
        doc.add_heading("Conversations", level=1)
        for chat in chats:
            doc.add_heading(f"Chat ID: {chat['chat_id']}", level=2)
            doc.add_paragraph(f"Timestamp: {chat['timestamp'].isoformat()}")
            doc.add_paragraph(f"User Message: {chat['message']}")
            doc.add_paragraph(f"AI Response: {chat['response']}")
            doc.add_paragraph("-" * 50)
   
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
   
    filename = f"meeting_{run_id}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.get("/{run_id}/download/csv")
async def download_conversations_csv(run_id: str, user_id: str = Depends(get_current_user)):
    meeting = await meetings.find_one({"run_id": run_id, "user_id": user_id})
    if not meeting:
        raise HTTPException(404, "Meeting not found")
   
    cursor = conversations.find({"run_id": run_id, "user_id": user_id}).sort("timestamp", 1)
    chats = await cursor.to_list(length=50)
   
    buffer = io.StringIO()
    writer = csv.writer(buffer, quoting=csv.QUOTE_MINIMAL)
    writer.writerow(["Meeting ID", "Meeting Name", "Date", "Language", "Status", "Type", "Content"])
   
    writer.writerow([
        meeting["run_id"],
        meeting.get("filename", "Untitled Meeting"),
        meeting.get("created_at", "").isoformat(),
        meeting.get("language", "en"),
        meeting.get("status", "unknown"),
        "Transcript",
        meeting.get("result", {}).get("final_output", "")
    ])
    writer.writerow(["", "", "", "", "", "Summary", meeting.get("result", {}).get("summary", "")])
    for point in meeting.get("result", {}).get("key_points", []):
        writer.writerow(["", "", "", "", "", "Key Point", point])
    for item in meeting.get("result", {}).get("action_items", []):
        writer.writerow(["", "", "", "", "", "Action Item", item])
   
    for chat in chats:
        writer.writerow([
            meeting["run_id"],
            meeting.get("filename", "Untitled Meeting"),
            chat["timestamp"].isoformat(),
            "",
            "",
            "User Message",
            chat["message"]
        ])
        writer.writerow([
            meeting["run_id"],
            meeting.get("filename", "Untitled Meeting"),
            chat["timestamp"].isoformat(),
            "",
            "",
            "AI Response",
            chat["response"]
        ])
   
    buffer.seek(0)
    filename = f"meeting_{run_id}.csv"
    return StreamingResponse(
        io.BytesIO(buffer.getvalue().encode('utf-8')),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )